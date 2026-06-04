from io import BytesIO
from pathlib import Path

from docx import Document
from openpyxl import Workbook
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parent
SAMPLES_DIR = ROOT / "samples"


def create_docx_sample() -> Path:
    path = SAMPLES_DIR / "retriflow-sample.docx"
    document = Document()
    document.add_heading("RetriFlow Tika Sample", level=1)
    document.add_paragraph("This DOCX sample validates heading, paragraph, table, and embedded-image extraction.")
    document.add_picture(_build_sample_figure_stream(), width=None, height=None)
    document.add_paragraph("图1 RetriFlow document parsing overview")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Field"
    table.rows[0].cells[1].text = "Value"
    table.rows[1].cells[0].text = "Owner"
    table.rows[1].cells[1].text = "RetriFlow"
    document.save(path)
    return path


def create_xlsx_sample() -> Path:
    path = SAMPLES_DIR / "retriflow-sample.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Metrics"
    sheet["A1"] = "Metric"
    sheet["B1"] = "Value"
    sheet["A2"] = "Latency"
    sheet["B2"] = "20%"
    sheet["A3"] = "Owner"
    sheet["B3"] = "RetriFlow"
    workbook.save(path)
    return path


def create_docx_image_only_sample() -> Path:
    path = SAMPLES_DIR / "retriflow-image-only.docx"
    document = Document()
    document.add_heading("RetriFlow Image-Only DOCX Sample", level=1)
    document.add_paragraph("This sample relies on OCR from the embedded image instead of a typed caption.")
    document.add_picture(_build_sample_figure_stream(), width=None, height=None)
    document.save(path)
    return path


def create_pdf_sample() -> Path:
    path = SAMPLES_DIR / "retriflow-sample.pdf"
    pdf = canvas.Canvas(str(path), pagesize=A4)
    pdf.setTitle("RetriFlow PDF Sample")
    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawString(72, 800, "RetriFlow PDF Sample")
    pdf.setFont("Helvetica", 12)
    pdf.drawString(72, 775, "This PDF sample validates paragraph, image unpacking, and page extraction.")
    image_reader = ImageReader(_build_sample_figure_stream())
    pdf.drawImage(image_reader, 72, 560, width=420, height=120, preserveAspectRatio=True, mask="auto")
    pdf.drawString(72, 540, "Figure 1 RetriFlow parsing workflow")
    pdf.showPage()
    pdf.setFont("Helvetica", 12)
    pdf.drawString(72, 800, "Second page paragraph for page tracking.")
    pdf.save()
    return path


def create_pdf_image_only_sample() -> Path:
    path = SAMPLES_DIR / "retriflow-image-only.pdf"
    pdf = canvas.Canvas(str(path), pagesize=A4)
    pdf.setTitle("RetriFlow Image-Only PDF Sample")
    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawString(72, 800, "RetriFlow Image-Only PDF Sample")
    pdf.setFont("Helvetica", 12)
    pdf.drawString(72, 775, "This sample relies on OCR from the embedded image instead of a typed caption.")
    image_reader = ImageReader(_build_sample_figure_stream())
    pdf.drawImage(image_reader, 72, 560, width=420, height=120, preserveAspectRatio=True, mask="auto")
    pdf.save()
    return path


def _build_sample_figure_stream() -> BytesIO:
    image = Image.new("RGB", (1200, 320), "white")
    draw = ImageDraw.Draw(image)
    font = _load_font()
    draw.rectangle((24, 24, 1176, 296), outline="black", width=4)
    draw.text((60, 90), "Figure 1 RetriFlow Parsing Workflow", fill="black", font=font)
    draw.text((60, 170), "OCR should detect this embedded image caption", fill="black", font=font)

    buffer = BytesIO()
    image.save(buffer, format="PNG")
    buffer.seek(0)
    return buffer


def _load_font() -> ImageFont.ImageFont:
    for name in ("DejaVuSans.ttf", "arial.ttf"):
        try:
            return ImageFont.truetype(name, 42)
        except OSError:
            continue
    return ImageFont.load_default()


def main() -> None:
    SAMPLES_DIR.mkdir(parents=True, exist_ok=True)
    created = [
        create_docx_sample(),
        create_docx_image_only_sample(),
        create_xlsx_sample(),
        create_pdf_sample(),
        create_pdf_image_only_sample(),
    ]
    for path in created:
        print(path)


if __name__ == "__main__":
    main()
